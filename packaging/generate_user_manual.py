# /// script
# requires-python = ">=3.13"
# dependencies = [
#   "python-docx>=1.2.0",
# ]
# ///

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SOFTWARE_NAME = "DataEX-G 数据清洗与空间分析软件"
VERSION = "V1.0"
DOCUMENT_TITLE = "用户操作手册"
PROJECT_ROOT = Path(__file__).resolve().parents[1]
IMAGE_DIRECTORY = PROJECT_ROOT / "docs" / "software-copyright" / "images"
OUTPUT_DIRECTORY = PROJECT_ROOT / "docs" / "software-copyright"
OUTPUT_PATH = OUTPUT_DIRECTORY / f"DataEX-G-{VERSION}-用户操作手册.docx"
ICON_PATH = PROJECT_ROOT / "icon" / "dg.png"

DARK_GREEN = "174C3C"
MID_GREEN = "2F6D57"
LIGHT_GREEN = "EAF3EE"
PALE_YELLOW = "FFF6DD"
LIGHT_GRAY = "F2F5F3"
TEXT_COLOR = RGBColor(37, 52, 45)


def set_run_font(run, name: str = "Microsoft YaHei", size: float = 10.5) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_COLOR


def shade_cell(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(
    cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=8)
    field_run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([field_begin, instruction, field_end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=8)


def configure_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 28, DARK_GREEN),
        ("Heading 1", 18, DARK_GREEN),
        ("Heading 2", 14, MID_GREEN),
        ("Heading 3", 11.5, MID_GREEN),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def configure_content_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(f"{SOFTWARE_NAME} {VERSION}　{DOCUMENT_TITLE}")
    set_run_font(header_run, size=8.5)
    header_run.font.color.rgb = RGBColor.from_string(MID_GREEN)

    section.footer.is_linked_to_previous = False
    add_page_field(section.footer.paragraphs[0])

    page_number = OxmlElement("w:pgNumType")
    page_number.set(qn("w:start"), "1")
    section._sectPr.append(page_number)


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(55)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ICON_PATH.is_file():
        run = paragraph.add_run()
        run.add_picture(str(ICON_PATH), width=Cm(5.2))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title_run = title.add_run(SOFTWARE_NAME)
    set_run_font(title_run, size=25)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DARK_GREEN)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(10)
    subtitle_run = subtitle.add_run(DOCUMENT_TITLE)
    set_run_font(subtitle_run, size=22)
    subtitle_run.bold = True

    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_before = Pt(18)
    version_run = version.add_run(VERSION)
    set_run_font(version_run, size=14)
    version_run.font.color.rgb = RGBColor.from_string(MID_GREEN)

    prepared = document.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.paragraph_format.space_before = Pt(155)
    prepared_run = prepared.add_run(
        f"编制日期：{datetime.now(timezone.utc).astimezone().date().isoformat()}"
    )
    set_run_font(prepared_run, size=10.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead)
        lead.bold = True
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_steps(document: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.65)
        number = paragraph.add_run(f"{index}. ")
        set_run_font(number)
        number.bold = True
        number.font.color.rgb = RGBColor.from_string(DARK_GREEN)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_note(document: Document, text: str, *, warning: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, PALE_YELLOW if warning else LIGHT_GREEN)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    paragraph = cell.paragraphs[0]
    label = paragraph.add_run("注意：" if warning else "说明：")
    set_run_font(label)
    label.bold = True
    text_run = paragraph.add_run(text)
    set_run_font(text_run)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(
    document: Document, headers: list[str], rows: list[list[str]], widths=None
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        shade_cell(header_cells[index], DARK_GREEN)
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, size=9)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                shade_cell(cell, LIGHT_GRAY)
        for column_index, value in enumerate(values):
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[column_index])
            paragraph = cells[column_index].paragraphs[0]
            run = paragraph.add_run(value)
            set_run_font(run, size=8.8)
        if widths:
            for column_index, width in enumerate(widths):
                cells[column_index].width = Cm(width)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(
    document: Document, filename: str, caption: str, width: float = 16.2
) -> None:
    image_path = IMAGE_DIRECTORY / filename
    if not image_path.is_file():
        raise RuntimeError(f"缺少手册截图：{image_path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9)
    caption_run.font.color.rgb = RGBColor.from_string(MID_GREEN)


def add_contents(document: Document) -> None:
    add_heading(document, "目录", 1)
    contents = [
        "1　软件概述",
        "2　运行与启动",
        "3　界面与数据文件",
        "4　数据清洗",
        "5　回归与相关性分析",
        "6　空间分析",
        "7　结果导出与文件命名",
        "8　常见问题",
        "9　使用注意事项",
    ]
    for item in contents:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_after = Pt(7)
        run = paragraph.add_run(item)
        set_run_font(run, size=11)
    add_note(
        document,
        "本文档中的界面图片使用项目示例数据生成，字段名称和结果仅用于说明操作。",
    )


def build_manual(document: Document) -> None:
    add_heading(document, "1　软件概述", 1)
    add_body(
        document,
        "DataEX-G 是运行于 Windows 的本地数据处理工具，用于检查、清洗、标准化、回归分析、相关性分析和空间分析。程序读取 CSV 或 XLSX 文件，计算完成后可导出 CSV 或 XLSX 结果。",
    )
    add_table(
        document,
        ["功能区", "主要功能"],
        [
            ["数据清洗", "空值检查与处理、文本问题检查、Min-Max 和 Z-score 标准化"],
            ["回归分析", "OLS、负二项回归、Logistic 回归及模型诊断"],
            ["相关性分析", "Pearson、Spearman 相关系数及显著性检验"],
            ["空间分析", "Moran's I、SLM/SAR、SEM、SDM、GWR 及空间诊断"],
            ["结果导出", "清洗结果、回归结果和空间结果导出为 CSV 或 XLSX"],
        ],
        widths=[3.2, 12.8],
    )
    add_note(
        document,
        "数据由本机程序处理。操作不会覆盖原文件，只有点击导出后才会生成新文件。",
    )

    document.add_page_break()
    add_heading(document, "2　运行与启动", 1)
    add_heading(document, "2.1　运行环境", 2)
    add_bullets(
        document,
        [
            "操作系统：Windows 10 或 Windows 11，64 位。",
            "浏览组件：Microsoft Edge WebView2 Runtime。",
            "数据文件：CSV 或 XLSX，单个文件不超过 20 MB。",
            "建议内存：8 GB 或以上；较大的空间模型需要更多内存和计算时间。",
        ],
    )
    add_heading(document, "2.2　启动程序", 2)
    add_steps(
        document,
        [
            "完整解压 DataEX-G 发布包。",
            "进入 DataEX-G 文件夹，双击 DataEX-G.exe。",
            "等待主界面出现。首次启动可能稍慢。",
            "使用结束后关闭主窗口，程序会同时停止本地分析服务。",
        ],
    )
    add_note(
        document,
        "不要单独复制 DataEX-G.exe。EXE 与同目录依赖文件必须一起保留。程序同一时间只允许启动一个实例。",
        warning=True,
    )
    add_body(
        document,
        r"启动失败时，先确认 WebView2 Runtime 已安装。运行日志位于 %LOCALAPPDATA%\DataEX-G\logs\desktop.log。",
    )

    document.add_page_break()
    add_heading(document, "3　界面与数据文件", 1)
    add_figure(document, "00-main-interface.png", "图 1　DataEX-G 主界面")
    add_body(
        document,
        "主界面顶部包含“数据清洗”“回归分析方法”“空间分析”三个窗口。切换窗口不会修改原始数据。每个分析窗口需要单独选择数据文件。",
    )
    add_heading(document, "3.1　数据文件要求", 2)
    add_bullets(
        document,
        [
            "CSV 支持 UTF-8、UTF-8 BOM 和 GB18030 编码。",
            "XLSX 默认读取第一个工作表。",
            "首行应为字段名称；每一行应代表一条观测记录。",
            "同一列尽量使用统一的数据类型和单位。",
            "回归与空间分析只提供可安全识别的纯数值列。",
            "文件预览最多显示前 100 行，分析和导出仍使用完整数据。",
        ],
    )
    add_note(
        document,
        "分析前保留原始文件副本。不要在原始文件上直接覆盖保存导出结果。",
        warning=True,
    )

    document.add_page_break()
    add_heading(document, "4　数据清洗", 1)
    add_heading(document, "4.1　选择并检查文件", 2)
    add_steps(
        document,
        [
            "进入“数据清洗”窗口。",
            "点击文件区域，或将 CSV/XLSX 文件拖入区域。",
            "确认文件名后点击“开始分析”。",
            "查看文件摘要和质量报告。",
        ],
    )
    add_figure(document, "01-file-selection.png", "图 2　选择数据文件")

    document.add_page_break()
    add_heading(document, "4.2　查看质量摘要", 2)
    add_figure(document, "02-quality-summary.png", "图 3　文件质量摘要")
    add_table(
        document,
        ["指标", "含义"],
        [
            ["数据行数", "不含字段标题的数据记录数量"],
            ["字段数量", "数据列数量"],
            ["空值单元格", "缺失值和空文本的总数"],
            ["重复数据行", "内容完全相同的重复行数量"],
            ["前后空格", "文本开头或结尾包含空白的单元格数量"],
            ["回车/换行", "文本内部包含回车或换行的单元格数量"],
        ],
        widths=[3.5, 12.5],
    )
    add_body(
        document,
        "质量报告按字段展示 pandas 类型、空值数量、识别出的内容类型和文本问题。“混合类型”表示同一列中出现不同内容类型，应在建模前统一。",
    )
    add_figure(document, "03-quality-report.png", "图 4　字段质量报告")

    document.add_page_break()
    add_heading(document, "4.3　设置空值处理", 2)
    add_table(
        document,
        ["选项", "处理结果", "使用提示"],
        [
            ["暂不处理", "保留全部空值", "用于先检查数据或暂不确定处理方式时"],
            [
                "删除空值行",
                "删除包含任意空值的整行",
                "可能减少样本量，导出前检查删除数量",
            ],
            ["摘出为空值表", "从主表移出，并生成独立空值表", "适合后续人工核对或补录"],
            ["使用 0 替换", "所有空值替换为数值 0", "仅在 0 具有明确业务含义时使用"],
        ],
        widths=[3.2, 5.4, 7.4],
    )
    add_note(
        document,
        "空值不一定等于 0。随意填充 0 会改变均值、相关性和回归结果。",
        warning=True,
    )

    add_heading(document, "4.4　设置文本处理", 2)
    add_bullets(
        document,
        [
            "清除文本前后空格：删除文本首尾空白，不删除文本中间的正常空格。",
            "清除回车和换行符：将连续回车/换行替换为一个空格。",
            "未勾选的文本问题只作提示，不会被修改。",
        ],
    )

    add_heading(document, "4.5　设置数据标准化", 2)
    add_table(
        document,
        ["方法", "计算结果", "适用情况"],
        [
            ["Min-Max", "数值缩放到 0—1", "需要统一量纲或固定范围时"],
            ["Z-score", "均值约为 0，标准差约为 1", "变量尺度差异较大或用于模型比较时"],
        ],
        widths=[3.2, 5.8, 7.0],
    )
    add_body(
        document,
        "选择标准化方法后，至少勾选一个数值列。常数列标准化后为 0。结果保留 6 位小数。",
    )
    add_figure(document, "04-cleaning-rules.png", "图 5　清洗规则与标准化设置")

    document.add_page_break()
    add_heading(document, "4.6　预览并导出清洗结果", 2)
    add_steps(
        document,
        [
            "完成规则设置后点击“生成清洗预览”。",
            "核对原始行数、清洗后行数、摘出行数和文本修改数量。",
            "检查主表预览；选择摘出空值表时，同时检查空值表预览。",
            "确认无误后导出主表 CSV、空值表 CSV 或 XLSX。",
        ],
    )
    add_figure(document, "05-cleaning-result.png", "图 6　清洗预览与导出")
    add_note(
        document,
        "XLSX 将清洗后主表写入 cleaned_data 工作表；存在摘出数据时，同时写入 missing_data 工作表。",
    )

    document.add_page_break()
    add_heading(document, "5　回归与相关性分析", 1)
    add_heading(document, "5.1　方法选择", 2)
    add_table(
        document,
        ["方法", "数据要求", "主要用途"],
        [
            ["OLS", "连续数值因变量；一个或多个数值自变量", "估计线性关系"],
            ["负二项回归", "非负整数计数因变量", "处理计数数据及过度离散"],
            ["Pearson", "两个数值变量", "检验线性相关"],
            ["Spearman", "两个可排序数值变量", "检验单调相关"],
            ["Logistic", "仅包含两类取值的因变量", "估计二分类结果"],
        ],
        widths=[3.0, 6.5, 6.5],
    )
    add_note(
        document,
        "相关性和统计显著性不代表因果关系。方法选择应与变量类型和研究设计一致。",
        warning=True,
    )

    add_heading(document, "5.2　运行分析", 2)
    add_steps(
        document,
        [
            "进入“回归分析方法”窗口并选择数据文件。",
            "选择分析方法。",
            "指定因变量 Y；回归方法勾选一个或多个自变量 X。",
            "Pearson 或 Spearman 分别选择变量 Y 和变量 X。",
            "点击“运行分析”。",
        ],
    )
    add_figure(document, "06-regression-settings.png", "图 7　OLS 方法与变量设置")

    document.add_page_break()
    add_heading(document, "5.3　查看分析结果", 2)
    add_bullets(
        document,
        [
            "样本信息：显示有效样本数和因缺失值排除的行数。",
            "诊断信息：显示是否收敛、条件数、最大 VIF、尺度比和警告。",
            "拟合指标：根据方法显示 R²、调整 R²、伪 R²、AIC 或 BIC。",
            "系数表：显示系数、标准误、统计量、p 值和 95% 置信区间。",
            "Logistic 结果额外显示优势比；负二项回归额外显示发生率比。",
        ],
    )
    add_figure(document, "07-regression-result.png", "图 8　OLS 分析结果与导出")
    add_note(
        document,
        "出现“模型未收敛”“模型推断不可靠”或严重共线性警告时，不应直接解释系数和 p 值。先检查变量、样本量、异常值和数据尺度。",
        warning=True,
    )

    document.add_page_break()
    add_heading(document, "6　空间分析", 1)
    add_heading(document, "6.1　数据与坐标要求", 2)
    add_bullets(
        document,
        [
            "每一行代表一个带坐标的空间观测对象。",
            "经纬度坐标选择“经纬度（WGS84）”，X 为经度，Y 为纬度。",
            "投影坐标选择“投影坐标 X/Y”，X、Y 应使用同一坐标参考系和单位。",
            "坐标列、因变量和自变量必须为数值列。",
            "缺失坐标或模型变量的行会被排除。",
        ],
    )
    add_note(
        document,
        "不要把经纬度和投影坐标混用。重复坐标会影响 KNN 和 GWR 结果。",
        warning=True,
    )

    add_heading(document, "6.2　空间方法", 2)
    add_table(
        document,
        ["方法", "作用"],
        [
            ["Moran's I", "检验一个变量的全局空间自相关"],
            ["SLM / SAR", "在模型中加入因变量空间滞后项 Wy"],
            ["SEM", "处理误差项中的空间相关结构"],
            ["SDM", "同时加入 Wy 和自变量空间滞后项 WX"],
            ["GWR", "估计随空间位置变化的局部回归系数"],
        ],
        widths=[3.5, 12.5],
    )

    document.add_page_break()
    add_heading(document, "6.3　设置并运行空间分析", 2)
    add_steps(
        document,
        [
            "进入“空间分析”窗口并选择数据文件。",
            "选择 Moran's I、SLM/SAR、SEM、SDM 或 GWR。",
            "选择坐标类型，并确认 X、Y 坐标列。常见字段名会自动识别。",
            "设置 K 近邻数。K 必须大于等于 1，且小于数据行数。",
            "选择空间自相关变量，或设置因变量 Y 和自变量 X。",
            "点击“运行空间分析”。",
        ],
    )
    add_figure(document, "08-spatial-settings.png", "图 9　Moran's I 空间分析设置")
    add_note(
        document,
        "K 值决定空间权重结构。正式分析应比较多个合理 K 值，并说明选择依据。",
        warning=True,
    )

    document.add_page_break()
    add_heading(document, "6.4　查看空间结果", 2)
    add_bullets(
        document,
        [
            "权重摘要：显示 KNN 邻居数、行标准化和连通分量。",
            "Moran's I：显示指数、期望值、Z 值和置换检验 p 值。",
            "空间回归：显示系数、拟合指标、ρ 或 λ，并提供模型诊断。",
            "SLM/SDM：显示直接效应、间接效应和总效应。",
            "残差 Moran：检查模型残差是否仍存在空间自相关。",
            "模型选择诊断：根据基础 OLS 的 LM 与稳健 LM 检验给出提示。",
            "GWR：显示带宽、拟合指标、局部系数分布和局部结果预览。",
        ],
    )
    add_figure(document, "09-spatial-result.png", "图 10　Moran's I 结果与导出")
    add_note(
        document,
        "模型选择提示仅供辅助。正式结论还需结合研究理论、空间权重敏感性、残差检验和模型拟合指标。",
        warning=True,
    )

    add_heading(document, "6.5　GWR 使用限制", 2)
    add_bullets(
        document,
        [
            "当前版本最多处理 5,000 条坐标和模型变量均完整的观测记录。",
            "样本量必须大于自动选择带宽所需的最低数量。",
            "存在重复坐标、严重共线性或局部样本不足时，模型可能无法估计。",
            "导出 GWR 时会重新拟合模型，并写入全部局部结果，耗时可能长于页面预览。",
            "局部 p_value_unadjusted 为未经多重检验校正的双侧近似 p 值。",
        ],
    )

    document.add_page_break()
    add_heading(document, "7　结果导出与文件命名", 1)
    add_table(
        document,
        ["结果类型", "文件名规则", "说明"],
        [
            ["清洗后主表", "原文件名-dataex.csv/.xlsx", "XLSX 可同时包含空值数据表"],
            ["摘出空值表 CSV", "原文件名-empty-dataex.csv", "仅在选择摘出空值表时生成"],
            [
                "回归/相关性结果",
                "原文件名-analysis-dataex.csv/.xlsx",
                "XLSX 包含概览和结果明细",
            ],
            [
                "空间分析结果",
                "原文件名-spatial-dataex.csv/.xlsx",
                "XLSX 包含模型、诊断和明细",
            ],
        ],
        widths=[3.4, 6.2, 6.4],
    )
    add_bullets(
        document,
        [
            "CSV 适合继续处理或导入其他软件。",
            "XLSX 适合同时保存概览、诊断和多张结果表。",
            "导出文件由系统下载到 Windows 当前默认下载位置。",
            "导出后应核对行数、字段、工作表和关键统计量。",
        ],
    )

    add_heading(document, "8　常见问题", 1)
    add_table(
        document,
        ["问题", "处理方法"],
        [
            ["程序无法启动", "确认完整解压程序目录并安装 WebView2；查看 desktop.log"],
            [
                "文件无法读取",
                "确认格式为 CSV/XLSX、文件非空、未超过 20 MB；CSV 改存为 UTF-8",
            ],
            ["没有可用数值列", "清理混合类型、空字符串和非数字字符，再重新导入"],
            ["运行按钮不可用", "检查方法、因变量、自变量、坐标列和 K 值是否完整"],
            [
                "模型未收敛",
                "减少共线变量，检查异常值、样本量和变量尺度，必要时标准化自变量",
            ],
            [
                "标准误或 p 值为空",
                "检查样本量、完全分离、奇异矩阵、非有限数值及共线性警告",
            ],
            ["空间模型失败", "检查坐标重复、K 值、连通性、缺失值和变量共线性"],
            ["导出耗时较长", "等待计算完成；GWR 导出需要重新拟合全部局部结果"],
        ],
        widths=[4.0, 12.0],
    )

    add_heading(document, "9　使用注意事项", 1)
    add_bullets(
        document,
        [
            "始终保留未经处理的原始数据。",
            "清洗预览确认无误后再导出，不要仅凭问题数量决定处理方式。",
            "删除空值行前检查样本损失；使用 0 填充前确认业务含义。",
            "标准化会改变变量尺度，导出后记录所用方法和字段。",
            "回归分析前确认因变量类型符合模型要求。",
            "不要把相关性、显著性或自动模型建议解释为因果关系。",
            "空间分析应记录坐标系统、K 值和空间权重设定。",
            "存在未收敛、非有限数值或严重共线性警告时，不直接使用模型结论。",
            "正式报告应结合研究设计、样本来源、模型假设和专业判断。",
        ],
    )
    add_note(
        document,
        "DataEX-G 提供数据处理和统计分析辅助，不替代数据质量审查、专业模型论证或业务决策责任。",
        warning=True,
    )


def main() -> None:
    OUTPUT_DIRECTORY.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    document.core_properties.title = f"{SOFTWARE_NAME} {VERSION} {DOCUMENT_TITLE}"
    document.core_properties.subject = "计算机软件著作权文档鉴别材料"
    document.core_properties.author = "DataEX-G"
    document.core_properties.keywords = (
        "DataEX-G, 用户操作手册, 数据清洗, 回归分析, 空间分析"
    )

    add_cover(document)
    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_content_section(content_section)
    add_contents(document)
    document.add_page_break()
    build_manual(document)
    document.save(OUTPUT_PATH)
    print(f"MANUAL={OUTPUT_PATH}")
    print(f"IMAGES={len(list(IMAGE_DIRECTORY.glob('*.png')))}")


if __name__ == "__main__":
    main()
