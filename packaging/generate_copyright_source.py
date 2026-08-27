# /// script
# requires-python = ">=3.13"
# dependencies = [
#   "python-docx>=1.2.0",
#   "reportlab>=4.4.0",
# ]
# ///

from __future__ import annotations

import unicodedata
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas

SOFTWARE_NAME = "DataEX-G 数据清洗与空间分析软件"
VERSION = "V1.0"
LINES_PER_PAGE = 50
PAGES_PER_SECTION = 30
DOCUMENT_PAGES = PAGES_PER_SECTION * 2
SELECTED_LINES_PER_SECTION = LINES_PER_PAGE * PAGES_PER_SECTION

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIRECTORY = PROJECT_ROOT / "docs" / "software-copyright"
OUTPUT_STEM = f"DataEX-G-{VERSION}-源代码文档"

SOURCE_FILES = [
    "backend/src/backend/__init__.py",
    "backend/src/backend/resources.py",
    "backend/src/backend/desktop.py",
    "backend/src/backend/quality.py",
    "backend/src/backend/cleaning.py",
    "backend/src/backend/diagnostics.py",
    "backend/src/backend/analysis.py",
    "backend/src/backend/spatial.py",
    "backend/src/backend/reporting.py",
    "backend/src/backend/main.py",
    "frontend/src/main.ts",
    "frontend/src/services/api.ts",
    "frontend/src/types/analysis.ts",
    "frontend/src/App.vue",
    "frontend/src/components/AnalysisWorkspace.vue",
    "frontend/src/components/SpatialWorkspace.vue",
]


@dataclass(frozen=True)
class SourceFileStat:
    path: str
    total_lines: int
    nonblank_lines: int


@dataclass(frozen=True)
class SourceLine:
    source_number: int
    text: str


def marker_for(path: str) -> str:
    suffix = Path(path).suffix
    if suffix == ".py":
        return f"# Source file: {path}"
    if suffix == ".vue":
        return f"<!-- Source file: {path} -->"
    return f"// Source file: {path}"


def collect_source() -> tuple[list[SourceLine], list[SourceFileStat]]:
    collected: list[str] = []
    stats: list[SourceFileStat] = []

    for relative_path in SOURCE_FILES:
        source_path = PROJECT_ROOT / relative_path
        raw_lines = source_path.read_text(encoding="utf-8").splitlines()
        code_lines = [line.expandtabs(4).rstrip() for line in raw_lines if line.strip()]
        stats.append(
            SourceFileStat(
                path=relative_path,
                total_lines=len(raw_lines),
                nonblank_lines=len(code_lines),
            )
        )
        collected.append(marker_for(relative_path))
        collected.extend(code_lines)

    numbered = [
        SourceLine(source_number=index, text=text)
        for index, text in enumerate(collected, start=1)
    ]
    return numbered, stats


def select_deposit_lines(source_lines: list[SourceLine]) -> list[SourceLine]:
    required = SELECTED_LINES_PER_SECTION * 2
    if len(source_lines) < required:
        raise RuntimeError(
            f"源代码只有 {len(source_lines)} 个非空展示行，不足以生成 {required} 行普通交存材料。"
        )
    return [
        *source_lines[:SELECTED_LINES_PER_SECTION],
        *source_lines[-SELECTED_LINES_PER_SECTION:],
    ]


def display_width(text: str) -> int:
    return sum(
        2 if unicodedata.east_asian_width(character) in {"W", "F"} else 1
        for character in text
    )


def render_line(line: SourceLine) -> str:
    return f"{line.source_number:04d}  {line.text}"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(8)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])

    suffix = paragraph.add_run(" 页")
    suffix.font.name = "Microsoft YaHei"
    suffix.font.size = Pt(8)


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(f"{SOFTWARE_NAME} {VERSION}　源程序鉴别材料")
    header_run.bold = True
    header_run.font.name = "Microsoft YaHei"
    header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header_run.font.size = Pt(9)

    add_page_number(section.footer.paragraphs[0])


def create_docx(lines: list[SourceLine], output_path: Path) -> None:
    document = Document()
    configure_docx(document)
    document.core_properties.title = f"{SOFTWARE_NAME} {VERSION} 源代码文档"
    document.core_properties.subject = "计算机软件著作权源程序鉴别材料"
    document.core_properties.author = "DataEX-G"

    for page_index in range(DOCUMENT_PAGES):
        if page_index:
            document.add_page_break()

        page_lines = lines[
            page_index * LINES_PER_PAGE : (page_index + 1) * LINES_PER_PAGE
        ]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(10.2)

        for line_index, source_line in enumerate(page_lines):
            text = render_line(source_line)
            size = max(4.8, min(6.5, 6.5 * 132 / max(display_width(text), 1)))
            run = paragraph.add_run(text)
            run.font.name = "Microsoft YaHei UI"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
            run.font.size = Pt(size)
            if line_index < LINES_PER_PAGE - 1:
                run.add_break(WD_BREAK.LINE)

    document.save(output_path)


def create_pdf(lines: list[SourceLine], output_path: Path) -> None:
    font_path = Path("C:/Windows/Fonts/simhei.ttf")
    if not font_path.is_file():
        raise RuntimeError("未找到 C:/Windows/Fonts/simhei.ttf，无法可靠输出中文 PDF。")

    font_name = "DataEXG-SimHei"
    pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    canvas = Canvas(str(output_path), pagesize=A4, pageCompression=1)
    canvas.setTitle(f"{SOFTWARE_NAME} {VERSION} 源代码文档")
    canvas.setAuthor("DataEX-G")
    page_width, page_height = A4
    left = 34
    right = page_width - 34
    available_width = right - left

    for page_index in range(DOCUMENT_PAGES):
        canvas.setFont(font_name, 9)
        canvas.drawCentredString(
            page_width / 2,
            page_height - 27,
            f"{SOFTWARE_NAME} {VERSION}　源程序鉴别材料",
        )

        page_lines = lines[
            page_index * LINES_PER_PAGE : (page_index + 1) * LINES_PER_PAGE
        ]
        y = page_height - 48
        for source_line in page_lines:
            text = render_line(source_line)
            base_size = 6.5
            measured_width = pdfmetrics.stringWidth(text, font_name, base_size)
            size = max(
                4.8,
                min(base_size, base_size * available_width / max(measured_width, 1)),
            )
            canvas.setFont(font_name, size)
            canvas.drawString(left, y, text)
            y -= 14.4

        canvas.setFont(font_name, 8)
        canvas.drawCentredString(page_width / 2, 18, f"第 {page_index + 1} 页")
        canvas.showPage()

    canvas.save()


def create_manifest(
    source_lines: list[SourceLine], stats: list[SourceFileStat], output_path: Path
) -> None:
    rear_start = len(source_lines) - SELECTED_LINES_PER_SECTION + 1
    rows = [
        f"# {SOFTWARE_NAME} {VERSION} 源代码文件清单",
        "",
        f"生成日期：{datetime.now(timezone.utc).astimezone().date().isoformat()}",
        "",
        "## 交存结构",
        "",
        f"- 文档页数：{DOCUMENT_PAGES} 页。",
        f"- 每页代码：{LINES_PER_PAGE} 行。",
        f"- 第 1—30 页：完整排列源程序第 1—{SELECTED_LINES_PER_SECTION} 行。",
        f"- 第 31—60 页：完整排列源程序第 {rear_start}—{len(source_lines)} 行。",
        f"- 排列后的自有源程序总计：{len(source_lines)} 个非空展示行（包括文件路径注释行）。",
        "",
        "## 文件顺序",
        "",
        "| 序号 | 源文件 | 原始行数 | 非空代码行 |",
        "|---:|---|---:|---:|",
    ]
    for index, stat in enumerate(stats, start=1):
        rows.append(
            f"| {index} | `{stat.path}` | {stat.total_lines} | {stat.nonblank_lines} |"
        )

    rows.extend(
        [
            "",
            "## 未纳入范围",
            "",
            "以下内容未作为申请人的自有核心源程序纳入鉴别材料：",
            "",
            "- `.venv/`、`node_modules/` 等第三方依赖。",
            "- `dist/`、`build/`、`frontend/dist/` 等生成产物。",
            "- `uv.lock`、`package-lock.json` 等依赖锁定文件。",
            "- 测试文件、示例数据和个人工作记录。",
            "",
            "## 提交前检查",
            "",
            "本清单用于材料整理和复核，不替代登记机构的最终要求。提交前请申请人逐页检查软件名称、版本号、代码真实性、权利归属、连续性和打印效果。",
            "",
        ]
    )
    output_path.write_text("\n".join(rows), encoding="utf-8")


def main() -> None:
    source_lines, stats = collect_source()
    deposit_lines = select_deposit_lines(source_lines)
    OUTPUT_DIRECTORY.mkdir(parents=True, exist_ok=True)

    docx_path = OUTPUT_DIRECTORY / f"{OUTPUT_STEM}.docx"
    pdf_path = OUTPUT_DIRECTORY / f"{OUTPUT_STEM}.pdf"
    manifest_path = OUTPUT_DIRECTORY / f"DataEX-G-{VERSION}-源代码文件清单.md"

    create_docx(deposit_lines, docx_path)
    create_pdf(deposit_lines, pdf_path)
    create_manifest(source_lines, stats, manifest_path)

    print(f"DOCX={docx_path}")
    print(f"PDF={pdf_path}")
    print(f"MANIFEST={manifest_path}")
    print(f"SOURCE_LINES={len(source_lines)}")
    print(f"DEPOSIT_LINES={len(deposit_lines)}")
    print(f"PAGES={DOCUMENT_PAGES}")


if __name__ == "__main__":
    main()
